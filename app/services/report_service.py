from pathlib import Path
from datetime import datetime
import json

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)


def get_original_text(item: dict) -> str:
    return (
        item.get("original_text")
        or item.get("highlight_text")
        or "-"
    )


def get_reason(item: dict) -> str:
    return (
        item.get("reason")
        or item.get("revision_reason")
        or item.get("issue_summary")
        or "-"
    )


def get_revision_detail(item: dict) -> str:
    return (
        item.get("revision_detail")
        or item.get("review_focus")
        or "-"
    )


def format_legal_basis(item: dict) -> str:
    legal_basis = item.get("legal_basis", [])

    if not legal_basis:
        evidence = item.get("evidence", [])
        legal_basis = evidence

    if not legal_basis:
        return "-"

    lines = []

    for basis in legal_basis:
        if not isinstance(basis, dict):
            continue

        source = basis.get("source", "-")
        page = basis.get("page", "-")
        article = basis.get("article", "")
        content = basis.get("content", "")

        text = f"{source}"
        if page not in [None, "", "-"]:
            text += f" p.{page}"
        if article:
            text += f" {article}"
        if content:
            text += f" - {content[:120]}"

        lines.append(text)

    return "\n".join(lines) if lines else "-"


def format_summary(summary: dict) -> str:
    if not isinstance(summary, dict):
        return str(summary)

    total_issues = (
        summary.get("total_issues")
        or summary.get("total_items")
        or 0
    )

    lines = [f"총 이슈 수: {total_issues}"]

    issue_counts = summary.get("issue_summary_counts", {})
    for issue, count in issue_counts.items():
        lines.append(f"- {issue}: {count}건")

    key_findings = summary.get("key_findings", [])
    if key_findings:
        lines.append("\n주요 발견사항:")
        for item in key_findings:
            lines.append(f"- {item}")

    return "\n".join(lines)


def create_txt_report(review: dict) -> Path:
    report_file_name = f"review_{review['review_id']}_report.txt"
    report_path = OUTPUT_DIR / report_file_name

    highlights_text = ""

    for idx, item in enumerate(review.get("highlights", []), start=1):
        highlights_text += f"""
{idx}. 수정 필요 사항
- 페이지: {item.get("page", "-")}
- 원문: {get_original_text(item)}
- 이슈: {item.get("issue_summary", "-")}
- 근거: {format_legal_basis(item)}
- 수정 제안: {item.get("suggested_text", "-")}
- 수정 이유: {get_reason(item)}
- 수정 설명: {get_revision_detail(item)}
"""

    revised_document = review.get("revised_document", "")

    report_path.write_text(
        f"""준법 리스크 검토 보고서

생성일시: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

문서 ID: {review.get("file_id")}
파일명: {review.get("file_name", "-")}
검토 언어: {review.get("language")}
규정 범위: {review.get("regulation_scope", "-")}

[요약]
{format_summary(review.get("summary", {}))}

[수정 필요 사항]
{highlights_text}

[수정안 반영 문안]
{revised_document or "-"}
""",
        encoding="utf-8",
    )

    return report_path


def create_docx_report(review: dict) -> Path:
    report_file_name = f"review_{review['review_id']}_report.docx"
    report_path = OUTPUT_DIR / report_file_name

    doc = Document()

    doc.add_heading("준법 리스크 검토 보고서", level=0)

    doc.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"문서 ID: {review.get('file_id')}")
    doc.add_paragraph(f"파일명: {review.get('file_name', '-')}")
    doc.add_paragraph(f"검토 언어: {review.get('language')}")
    doc.add_paragraph(f"규정 범위: {review.get('regulation_scope', '-')}")

    doc.add_heading("1. 검토 요약", level=1)
    doc.add_paragraph(format_summary(review.get("summary", {})))

    doc.add_heading("2. 수정 필요 사항", level=1)

    highlights = review.get("highlights", [])

    if not highlights:
        doc.add_paragraph("수정 필요 사항이 없습니다.")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "페이지"
        header_cells[1].text = "원문"
        header_cells[2].text = "이슈"
        header_cells[3].text = "근거"
        header_cells[4].text = "수정 제안"

        for item in highlights:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("page", "-"))
            row_cells[1].text = get_original_text(item)
            row_cells[2].text = item.get("issue_summary", "-")
            row_cells[3].text = format_legal_basis(item)
            row_cells[4].text = item.get("suggested_text", "-")

    doc.add_heading("3. 수정안 반영 문안", level=1)
    doc.add_paragraph(review.get("revised_document", "-") or "-")

    doc.add_heading("4. 종합 의견", level=1)
    doc.add_paragraph(
        "본 보고서는 업로드된 문서의 준법 리스크를 자동 검토한 결과입니다. "
        "실제 법률 판단이 필요한 경우 담당자의 추가 검토가 필요합니다."
    )

    doc.save(report_path)

    return report_path


def create_pdf_report(review: dict) -> Path:
    report_file_name = f"review_{review['review_id']}_report.pdf"
    report_path = OUTPUT_DIR / report_file_name

    c = canvas.Canvas(str(report_path), pagesize=A4)
    width, height = A4

    x = 50
    y = height - 50

    c.setFont("Helvetica-Bold", 16)
    c.drawString(x, y, "Compliance Risk Review Report")

    y -= 40
    c.setFont("Helvetica", 10)
    c.drawString(x, y, f"Created At: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    y -= 25
    c.drawString(x, y, f"Review ID: {review.get('review_id')}")
    y -= 20
    c.drawString(x, y, f"File ID: {review.get('file_id')}")
    y -= 20
    c.drawString(x, y, f"File Name: {review.get('file_name', '-')}")
    y -= 20
    c.drawString(x, y, f"Language: {review.get('language')}")
    y -= 20
    c.drawString(x, y, f"Regulation Scope: {review.get('regulation_scope', '-')}")

    y -= 35
    c.setFont("Helvetica-Bold", 13)
    c.drawString(x, y, "Summary")

    y -= 20
    c.setFont("Helvetica", 10)
    summary_text = format_summary(review.get("summary", {})).replace("\n", " / ")
    c.drawString(x, y, summary_text[:100])

    y -= 35
    c.setFont("Helvetica-Bold", 13)
    c.drawString(x, y, "Highlights")

    y -= 25
    c.setFont("Helvetica", 9)

    for idx, item in enumerate(review.get("highlights", []), start=1):
        if y < 120:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 9)

        c.drawString(x, y, f"{idx}. Page: {item.get('page', '-')}")
        y -= 15
        c.drawString(x, y, f"Original: {get_original_text(item)[:90]}")
        y -= 15
        c.drawString(x, y, f"Issue: {item.get('issue_summary', '-')[:90]}")
        y -= 15
        c.drawString(x, y, f"Suggestion: {item.get('suggested_text', '-')[:90]}")
        y -= 25

    c.save()

    return report_path


def create_reports(review: dict) -> dict:
    txt_path = create_txt_report(review)
    docx_path = create_docx_report(review)
    pdf_path = create_pdf_report(review)

    return {
        "txt": txt_path.name,
        "docx": docx_path.name,
        "pdf": pdf_path.name,
    }