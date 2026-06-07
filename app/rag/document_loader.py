from pathlib import Path
from typing import List

import fitz

from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument

def read_pdf(file_path: str) -> List[LCDocument]:
    docs = []
    path = Path(file_path)

    pdf = fitz.open(file_path)

    for page_num, page in enumerate(pdf, start=1):
        text = page.get_text()
        if text.strip():
            docs.append(
                LCDocument(
                    page_content=text,
                    metadata={
                        "source": path.name,
                        "page": page_num,
                        "type": "pdf",
                    },
                )
            )

    return docs


def read_docx(file_path: str) -> List[LCDocument]:
    path = Path(file_path)
    doc = DocxDocument(file_path)

    texts = []

    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                texts.append(" | ".join(row_text))

    return [
        LCDocument(
            page_content="\n".join(texts),
            metadata={
                "source": path.name,
                "page": None,
                "type": "docx",
            },
        )
    ]



def read_txt(file_path: str) -> List[LCDocument]:
    path = Path(file_path)
    file_bytes = Path(file_path).read_bytes()

    text = None
    for enc in ["utf-8-sig", "utf-8", "cp949", "euc-kr"]:
        try:
            text = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        raise ValueError(f"{path.name} 파일 인코딩을 읽을 수 없습니다.")

    return [
        LCDocument(
            page_content=text,
            metadata={
                "source": path.name,
                "page": None,
                "type": "txt",
            },
        )
    ]


def load_file(file_path: str) -> List[LCDocument]:
    ext = Path(file_path).suffix.lower().replace(".", "")

    if ext == "pdf":
        return read_pdf(file_path)
    if ext == "docx":
        return read_docx(file_path)
    if ext == "txt":
        return read_txt(file_path)

    raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")

def load_target_document(file_path: str) -> str:
    docs = load_file(file_path)

    return "\n\n".join(
        doc.page_content
        for doc in docs
        if doc.page_content and doc.page_content.strip()
    )